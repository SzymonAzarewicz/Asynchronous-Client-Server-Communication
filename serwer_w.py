import socket
import threading
from PIL import Image
import json
import base64
import io
import os
import datetime
from docx import Document

class Client:
    def __init__(self, socket, address, name=None):
        self.socket = socket
        self.address = address
        self.name = name or f"Klient_{address[1]}"

clients = []

# Folder dla odbieranych plików
DOCX_FOLDER = "odebrane_pliki"
if not os.path.exists(DOCX_FOLDER):
    os.makedirs(DOCX_FOLDER)

def broadcast(message, sender=None):
    """Wysyła wiadomość do wszystkich klientów oprócz nadawcy"""
    for client in clients:
        if client != sender:
            try:
                client.socket.send(message)
            except:
                remove_client(client)

def remove_client(client):
    """Usuwa klienta z listy"""
    if client in clients:
        clients.remove(client)
        print(f"Klient {client.name} rozłączony.")

def image_to_ascii(image_data, width):
    """Convert an image to ASCII art with dynamic sizing"""
    try:
        # Decode base64 image data
        image_bytes = base64.b64decode(image_data)
        image = Image.open(io.BytesIO(image_bytes))
        
        # Get original dimensions
        orig_width, orig_height = image.size
        
        # Calculate proportional height
        aspect_ratio = orig_height / orig_width
        height = int(aspect_ratio * width * 0.5)  # Multiply by 0.5 to account for character aspect ratio
        
        # Resize image
        image = image.resize((width, height))
        image = image.convert('L')
        
        # Define ASCII characters from dark to light
        chars = '@%#*+=-:. '
        
        # Convert pixels to ASCII
        ascii_str = ""
        for y in range(height):
            for x in range(width):
                pixel_value = image.getpixel((x, y))
                char_idx = int(pixel_value * (len(chars)-1) / 255)
                ascii_str += chars[char_idx]
            ascii_str += "\n"
        
        return ascii_str
    except Exception as e:
        return f"Błąd konwersji: {str(e)}"

def extract_docx_text(file_bytes):
    """Ekstrahuje tekst z pliku DOCX"""
    try:
        # Utworzenie obiektu dokumentu z danych binarnych
        doc_io = io.BytesIO(file_bytes)
        document = Document(doc_io)
        
        # Ekstrakcja tekstu z dokumentu
        full_text = []
        
        # Dodanie tytułu dokumentu, jeśli dostępny
        if document.core_properties.title:
            full_text.append(f"Tytuł: {document.core_properties.title}")
            full_text.append("=" * 40)
        
        # Ekstrakcja tekstu z paragrafów
        for para in document.paragraphs:
            if para.text:
                full_text.append(para.text)
        
        # Ekstrakcja tekstu z tabel
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                if any(row_text):
                    full_text.append(" | ".join(row_text))
        
        # Połączenie całości w jeden string
        return "\n".join(full_text)
    except Exception as e:
        return f"Błąd ekstrakcji tekstu: {str(e)}"

def save_docx_file(file_data, file_name, client_name):
    """Zapisuje otrzymany plik .docx i ekstrahuje jego zawartość"""
    try:
        # Dekodowanie danych pliku
        file_bytes = base64.b64decode(file_data)
        
        # Tworzenie nazwy pliku z timestampem
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_client_name = ''.join(c for c in client_name if c.isalnum() or c in [' ', '_']).strip().replace(' ', '_')
        
        # Tworzenie nazwy pliku
        base_name, ext = os.path.splitext(file_name)
        new_filename = f"{base_name}_{safe_client_name}_{timestamp}{ext}"
        file_path = os.path.join(DOCX_FOLDER, new_filename)
        
        # Zapisywanie pliku
        with open(file_path, 'wb') as f:
            f.write(file_bytes)
        
        # Ekstrahowanie zawartości dokumentu
        document_text = extract_docx_text(file_bytes)
            
        return {
            'path': file_path,
            'text': document_text
        }
    except Exception as e:
        print(f"Błąd zapisywania pliku: {str(e)}")
        return None

def handle_client(client):
    while True:
        try:
            data = client.socket.recv(1024000)  # Increased buffer size for image data
            if not data:
                break
            
            try:
                request = json.loads(data.decode())
                client_name = request.get('client_name', client.name)
                
                if request.get('type') == 'image_to_ascii' and 'image_data' in request:
                    # Process image to ASCII conversion
                    width = request.get('width', 60)
                    ascii_art = image_to_ascii(request['image_data'], width)
                    
                    # Print ASCII art on server
                    print(f"\nOtrzymano żądanie konwersji obrazu od {client_name}. Wynik:")
                    print(ascii_art)
                    
                    # Send the ASCII art back to client
                    response = json.dumps({'type': 'ascii_response', 'data': ascii_art})
                    client.socket.send(response.encode())
                
                elif request.get('type') == 'docx_file' and 'file_data' in request and 'file_name' in request:
                    # Obsługa pliku .docx
                    file_data = request.get('file_data')
                    file_name = request.get('file_name')
                    
                    print(f"\nOtrzymano plik DOCX od {client_name}: {file_name}")
                    
                    # Zapisz plik i ekstrahuj zawartość
                    result = save_docx_file(file_data, file_name, client_name)
                    
                    if result:
                        saved_path = result['path']
                        document_text = result['text']
                        
                        # Wyświetl zawartość pliku
                        print("\n===== ZAWARTOŚĆ DOKUMENTU =====")
                        print(document_text)
                        print("==============================\n")
                        
                        print(f"Zapisano plik: {saved_path}")
                        
                        # Wyślij potwierdzenie do klienta
                        response = json.dumps({
                            'type': 'docx_response', 
                            'message': f"Plik {file_name} został pomyślnie odebrany i zapisany na serwerze."
                        })
                        client.socket.send(response.encode())
                        
                        # Powiadom innych klientów o otrzymaniu pliku z treścią
                        notification = json.dumps({
                            'type': 'text',
                            'message': f"Użytkownik {client_name} przesłał plik {file_name} do serwera."
                        })
                        broadcast(notification.encode(), client)
                        
                        # Wyślij treść dokumentu do wszystkich
                        document_notification = json.dumps({
                            'type': 'text',
                            'message': f"Zawartość dokumentu {file_name} od {client_name}:\n\n{document_text}"
                        })
                        broadcast(document_notification.encode(), client)
                        
                    else:
                        # Wyślij informację o błędzie
                        response = json.dumps({
                            'type': 'docx_response', 
                            'message': f"Wystąpił błąd podczas zapisywania pliku {file_name} na serwerze."
                        })
                        client.socket.send(response.encode())
                
                elif request.get('type') == 'text':
                    # Handle text message
                    message = request.get('message', '')
                    print(f"\nOdebrano od {client_name}: {message}")
                    response = json.dumps({
                        'type': 'text',
                        'message': f"Wiadomość od {client_name}: {message}"
                    })
                    broadcast(response.encode(), client)
                    
                else:
                    # Normal echo behavior
                    print(f"\nOdebrano od {client_name}: {data.decode()}")
                    client.socket.send(data)
                    
            except json.JSONDecodeError:
                # Normal text message handling
                print(f"\nOdebrano od {client.name}: {data.decode()}")
                client.socket.send(data)
                
        except Exception as e:
            print(f"Błąd obsługi klienta {client.name}: {str(e)}")
            break
            
    remove_client(client)
    client.socket.close()

def main():
    server = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    server.bind(('localhost', 8888))
    server.listen(5)
    print("Serwer nasłuchuje...")
    print(f"Pliki DOCX będą zapisywane w folderze: {os.path.abspath(DOCX_FOLDER)}")

    while True:
        client_socket, addr = server.accept()
        client = Client(client_socket, addr)
        clients.append(client)
        print(f"Połączono z {client.name} ({addr})")
        client_thread = threading.Thread(target=handle_client, args=(client,))
        client_thread.start()

if __name__ == "__main__":
    main()